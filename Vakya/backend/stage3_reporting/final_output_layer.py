import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_docx_report(analyzed_clauses: list[dict], executive_summary: dict, output_path: str):
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Vakya AI Contract Audit Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Executive Summary Section
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(executive_summary.get('executive_summary', 'N/A'))
    
    doc.add_heading('Overall Risk Assessment', level=2)
    doc.add_paragraph(executive_summary.get('overall_risk_assessment', 'N/A'))
    
    doc.add_heading('Key Recommendations', level=2)
    recommendations = executive_summary.get('key_recommendations', [])
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')
        
    doc.add_page_break()
    
    # Clause-by-clause detailed analysis
    doc.add_heading('Clause-by-Clause Detailed Analysis', level=1)
    
    for ac in analyzed_clauses:
        clause_id = ac.get("clause_id", "?")
        category = ac.get("clause_info", {}).get("category", "Unknown")
        risk_level = ac.get("risk_info", {}).get("risk_level", "Unknown")
        
        h2 = doc.add_heading(f'Clause {clause_id}: {category} [Risk: {risk_level}]', level=2)
        
        # Original Text
        doc.add_heading('Original Clause:', level=3)
        doc.add_paragraph(ac.get("text", ""))
        
        # Risk & Compliance
        doc.add_heading('Risk & Compliance Analysis:', level=3)
        risk_exp = ac.get("risk_info", {}).get("risk_explanation", "")
        comp_issues = ac.get("compliance_info", {}).get("compliance_issues", "")
        unfair = ac.get("risk_info", {}).get("unfair_one_sided_terms", [])
        
        p_risk = doc.add_paragraph()
        p_risk.add_run('Explanation: ').bold = True
        p_risk.add_run(f'{risk_exp}\n')
        p_risk.add_run('Compliance Issues: ').bold = True
        p_risk.add_run(f'{comp_issues}\n')
        p_risk.add_run('Unfair Terms: ').bold = True
        p_risk.add_run(f'{", ".join(unfair) if unfair else "None"}')
        
        # Negotiation Strategy
        doc.add_heading('Negotiation Strategy:', level=3)
        improved = ac.get("negotiation_info", {}).get("improved_wording", "No change needed")
        counter_terms = ac.get("negotiation_info", {}).get("counter_terms", [])
        
        p_neg = doc.add_paragraph()
        p_neg.add_run('Suggested Redline: ').bold = True
        p_neg.add_run(f'{improved}')
        
        if counter_terms:
            doc.add_paragraph("Counter Terms:", style='Body Text')
            for term in counter_terms:
                doc.add_paragraph(term, style='List Bullet')
                
        doc.add_paragraph('_' * 50) # separator
        
    doc.save(output_path)
    return output_path
